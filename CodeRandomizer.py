# Program that analyzes the code and replaces it with gibberish symbols
# The code is used to modify C / C++ code
# Code must first be in the same folder in docx format
# Install docx reader:
# pip install python-docx
# By: Matas Sabaliauskas
# Email: matas.sabal@gmail.com


# Write down the name of your code docx file
FILENAME = 'problem2.docx'

# Write down how the new file with generated code will be named
NEWFILENAME = FILENAME + '_new.docx'

#The gibberish symbols are:
randomkeyword = 'l'  # all of code will be multiples of this
randomkeyword2 = 'I' # for semi colon
counter = 0;


from collections import Counter
import docx


#################################################
#DOCX READER
def ReadingTextDocuments(filename):
    doc = docx.Document(FILENAME)
    completedText = []
    for paragraph in doc.paragraphs:
        completedText.append(paragraph.text)
    return '\n' .join(completedText)

#to read the entire file:
#print(ReadingTextDocuments(FILENAME))

#################################################
#################################################
#Program finds all unique, separate words

data_set_docx = ReadingTextDocuments(FILENAME)
split_it_docx = data_set_docx.split()   # split() returns list of all the words in the string
Counter_docx = Counter(split_it_docx)   # Pass the split_it list to instance of Counter class.
most_frequent_docx = Counter_docx.most_common(100000)
print("The most frequent keywords in DOCX were:")
print(split_it_docx)


################################################
################################################

# Creates a document and adds modified text
doc = docx.Document()
paragraph1 = doc.add_paragraph()
paragraph2 = doc.add_paragraph()
paragraph3 = doc.add_paragraph('\n\n\n')
paragraph3.add_run('#include <iostream>' + '\n') #optional
paragraph3.add_run('#include <string>' + '\n') #optional

split_it_docx_mod = split_it_docx
while(counter < 1): #arbitrary while loop which doesnt work. crashes without it
#while(split_it_docx_mod != ''):
    for (split_it_docx_mod[counter]) in split_it_docx_mod:
        try:
            if (";" in split_it_docx_mod[counter]):                         #handles semicolon cases
                new = split_it_docx_mod[counter].replace (';', '')          #makes the string without semi colon
                paragraph2.add_run((randomkeyword*(counter+1)) + ' ')       #adds the eeee bit to the top list

                #paragraph2.add_run((randomkeyword2) + '\n')                # adds the x bit to the top list
                paragraph2.add_run(randomkeyword2 + '\n')                   # adds the x bit to the top list

                paragraph1 = doc.add_paragraph('#define ' + (randomkeyword*(counter+1)) + ' ')  #writes #define eeee word
                paragraph1.add_run(new)

                paragraph3.add_run('#define ' + (randomkeyword2) + ' ;' + '\n')        #prints only semicolon
                counter += 1

            elif (split_it_docx_mod[counter] == '') or (split_it_docx_mod[counter] == '#include') or (
                    split_it_docx_mod[counter] == '<iostream>') or (split_it_docx_mod[counter] == '<string>'):
                #paragraph1 = doc.add_paragraph(split_it_docx_mod[counter])
                #paragraph1.add_run(split_it_docx_mod[counter])
                counter += 1

            else:
                old = split_it_docx_mod[counter]
                print('Old keyword: ' + split_it_docx_mod[counter])

                if(split_it_docx_mod[counter] == '{'):
                    paragraph2.add_run('\n')

                split_it_docx_mod[counter] = randomkeyword * (counter + 1)
                newtext = split_it_docx_mod[counter]
                print('New keyword: ' + newtext)
                paragraph1 = doc.add_paragraph('#define ' + newtext + ' ' + old)

                paragraph2.add_run(newtext + ' ')
                counter += 1
                doc.save(NEWFILENAME)
        except IndexError:
            None

###########################################